# file_validator.py
import os
import re
import PyPDF2
import docx
from PIL import Image
import pytesseract
from datetime import datetime
import json

# Настройка пути к Tesseract (раскомментируйте для Windows)
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR'


def extract_text_from_pdf(file_path):
    """Извлечение текста из PDF с поддержкой сканов"""
    text = ""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text += page_text + "\n"
                else:
                    # Если текст не извлекся (возможно, это скан), пробуем OCR
                    print(f"Страница {page_num + 1} не содержит текста, пробуем OCR...")
        return text
    except Exception as e:
        print(f"Ошибка чтения PDF: {e}")
        return ""


def extract_text_from_docx(file_path):
    """Извлечение текста из DOCX"""
    text = ""
    try:
        doc = docx.Document(file_path)
        for paragraph in doc.paragraphs:
            if paragraph.text:
                text += paragraph.text + "\n"
        # Также извлекаем текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text += cell.text + "\n"
    except Exception as e:
        print(f"Ошибка чтения DOCX: {e}")
    return text


def extract_text_from_image(file_path):
    """Извлечение текста из изображения с помощью OCR"""
    text = ""
    try:
        # Открываем изображение
        image = Image.open(file_path)
        # Конвертируем в RGB если нужно
        if image.mode == 'RGBA':
            image = image.convert('RGB')
        # Увеличиваем контраст для лучшего распознавания
        from PIL import ImageEnhance
        enhancer = ImageEnhance.Contrast(image)
        image = enhancer.enhance(1.5)
        # Распознаем русский и английский текст
        text = pytesseract.image_to_string(image, lang='rus+eng')
        print(f"OCR распознано {len(text)} символов")
    except Exception as e:
        print(f"Ошибка OCR: {e}")
        # Пробуем альтернативный путь
        try:
            text = pytesseract.image_to_string(Image.open(file_path), lang='rus')
        except:
            pass
    return text


def extract_text_from_file(file_path):
    """Универсальное извлечение текста из файла"""
    ext = os.path.splitext(file_path)[1].lower()

    print(f"Обработка файла: {file_path}, расширение: {ext}")

    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext in ['.docx', '.doc']:
        return extract_text_from_docx(file_path)
    elif ext in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
        return extract_text_from_image(file_path)
    else:
        # Для .txt файлов
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except:
            return ""


def validate_consent_file(file_path, applicant_data):
    """Проверка согласия на ОПД"""
    text = extract_text_from_file(file_path)

    print(f"Извлеченный текст (первые 200 символов): {text[:200] if text else 'ПУСТО'}")

    if not text or len(text.strip()) < 50:
        return {
            'valid': False,
            'errors': ['Не удалось извлечь текст из файла. Убедитесь, что файл не защищен паролем и содержит текст.'],
            'warnings': []
        }

    errors = []
    warnings = []

    # Приводим к нижнему регистру для поиска
    text_lower = text.lower()

    # Проверка ФИО
    full_name = applicant_data.get('full_name', '')
    if full_name:
        name_parts = full_name.lower().split()
        name_found = 0
        for part in name_parts:
            if len(part) > 2 and part in text_lower:
                name_found += 1

        if name_found < 2:
            warnings.append(f"ФИО '{full_name}' не найдено полностью в согласии")

    # Проверка паспорта
    passport_series = applicant_data.get('passport_series', '')
    passport_number = applicant_data.get('passport_number', '')

    if passport_series and passport_number:
        passport_full = f"{passport_series} {passport_number}"
        passport_full_no_space = f"{passport_series}{passport_number}"

        if passport_full not in text and passport_full_no_space not in text:
            warnings.append(f"Паспортные данные не найдены в согласии")

    # Проверка ключевых фраз
    consent_phrases = ['согласие на обработку', 'персональных данных', 'даю согласие']
    has_consent = any(phrase in text_lower for phrase in consent_phrases)

    if not has_consent:
        errors.append("В документе отсутствуют ключевые фразы о согласии на обработку ПД")

    return {
        'valid': len(errors) == 0,
        'errors': errors,
        'warnings': warnings
    }


def validate_education_certificate(file_path, applicant_data):
    """Проверка справки об обучении"""
    text = extract_text_from_file(file_path)

    if not text or len(text.strip()) < 50:
        return {
            'valid': False,
            'errors': ['Не удалось извлечь текст из файла справки'],
            'warnings': []
        }

    errors = []
    warnings = []

    text_lower = text.lower()

    # Проверка ФИО
    full_name = applicant_data.get('full_name', '')
    if full_name:
        name_parts = full_name.lower().split()
        name_found = 0
        for part in name_parts:
            if len(part) > 2 and part in text_lower:
                name_found += 1

        if name_found < 1:
            warnings.append(f"ФИО не найдено в справке")

    # Проверка года
    current_year = datetime.now().year
    year_pattern = r'\b(20[2-9][0-9])\b'
    years = re.findall(year_pattern, text)

    found_valid = False
    for year in years:
        year_int = int(year)
        if year_int >= current_year - 1:
            found_valid = True
            break

    if not found_valid:
        warnings.append(f"Не найден актуальный год выдачи справки")

    return {
        'valid': len(errors) == 0,
        'errors': errors,
        'warnings': warnings
    }


def validate_all_documents(consent_path, education_path, applicant_data):
    """Комплексная проверка документов"""
    consent_result = validate_consent_file(consent_path, applicant_data)
    education_result = validate_education_certificate(education_path, applicant_data)

    return {
        'overall_valid': consent_result['valid'] and education_result['valid'],
        'consent': consent_result,
        'education': education_result,
        'total_errors': len(consent_result['errors']) + len(education_result['errors']),
        'total_warnings': len(consent_result['warnings']) + len(education_result['warnings'])
    }